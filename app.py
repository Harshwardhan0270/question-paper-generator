import streamlit as st
from langchain_groq import ChatGroq  # type: ignore
from langchain_core.output_parsers import StrOutputParser  # type: ignore
from langchain_core.prompts import ChatPromptTemplate  # type: ignore
import os
import docx  # type: ignore
from io import BytesIO
import time
import os
from dotenv import load_dotenv

load_dotenv()
# Set the Streamlit page configuration
st.set_page_config(
    page_title="Question Generator",
    page_icon="📘",
    layout="wide"
)

# Initialize the language model (make sure to keep your API key secure)


api_key = os.getenv("GROQ_API_KEY")
if not api_key:
    raise ValueError("GROQ_API_KEY environment variable is not set. Please set it before running the app.")

llm = ChatGroq(
    api_key=api_key,
    model="llama3-70b-8192"  # Or any valid Groq model
)


# Function to generate MCQs
def generate_mcq_questions(subject_name, syllabus, num_mcq, bl_level):
    prompt_template = """
    You are an expert in Bloom's Taxonomy and question generation. Based on the inputs below, generate {num_mcq} multiple-choice questions (MCQs) using the following format.

    Use the following short codes for Bloom’s levels:

    BL1: Remember  
    BL2: Understand  
    BL3: Apply  
    BL4: Analyze  
    BL5: Evaluate  
    BL6: Create

    {level_instruction}

    Format:
    Q1. What is the capital of France? [BL1]  
    (a) Berlin (b) Madrid (c) Paris (d) Rome

    Do NOT provide answers. Only questions.

    Subject: {subject_name}  
    Syllabus: {syllabus}
    """
    level_instruction = f"Generate ONLY {bl_level} level questions." if bl_level != "Random (All Levels)" else "Generate questions across all Bloom's levels."
    prompt = prompt_template.format(subject_name=subject_name, syllabus=syllabus, num_mcq=num_mcq, level_instruction=level_instruction)
    chain = (ChatPromptTemplate.from_template(prompt) | llm | StrOutputParser())
    return chain.invoke({})

# Function to generate short answer questions
def generate_short_questions(subject_name, syllabus, num_short, bl_level):
    prompt_template = """
    You are an expert in Bloom's Taxonomy and question generation. Generate {num_short} short answer questions using Bloom's short codes:

    BL1: Remember  
    BL2: Understand  
    BL3: Apply  
    BL4: Analyze  
    BL5: Evaluate  
    BL6: Create

    {level_instruction}

    Format:
    Q1. Explain the process of normalization. [BL2]

    Do NOT provide answers. Only questions.

    Subject: {subject_name}  
    Syllabus: {syllabus}
    """
    level_instruction = f"Generate ONLY {bl_level} level questions." if bl_level != "Random (All Levels)" else "Generate questions across all Bloom's levels."
    prompt = prompt_template.format(subject_name=subject_name, syllabus=syllabus, num_short=num_short, level_instruction=level_instruction)
    chain = (ChatPromptTemplate.from_template(prompt) | llm | StrOutputParser())
    return chain.invoke({})

# Function to generate long answer questions
def generate_long_questions(subject_name, syllabus, num_long, bl_level):
    prompt_template = """
    You are an expert in Bloom's Taxonomy and question generation. Generate {num_long} long answer questions using Bloom's short codes:

    BL1: Remember  
    BL2: Understand  
    BL3: Apply  
    BL4: Analyze  
    BL5: Evaluate  
    BL6: Create

    {level_instruction}

    Format:
    Q1. Design and implement a compiler for a toy language. [BL6]

    Do NOT provide answers. Only questions.

    Subject: {subject_name}  
    Syllabus: {syllabus}
    """
    level_instruction = f"Generate ONLY {bl_level} level questions." if bl_level != "Random (All Levels)" else "Generate questions across all Bloom's levels."
    prompt = prompt_template.format(subject_name=subject_name, syllabus=syllabus, num_long=num_long, level_instruction=level_instruction)
    chain = (ChatPromptTemplate.from_template(prompt) | llm | StrOutputParser())
    return chain.invoke({})

# Initialize session state
if "mcq_questions" not in st.session_state:
    st.session_state.mcq_questions = ""
if "short_questions" not in st.session_state:
    st.session_state.short_questions = ""
if "long_questions" not in st.session_state:
    st.session_state.long_questions = ""

# Title and header
st.title("Test Paper Generator using :red[Bloom's Taxonomy]📚")

# Sidebar with a custom background and title
st.sidebar.image("assets/compressed_becd011b419db54d4ea278a2d0425d7b.png")
st.sidebar.title("Question Generator")
st.sidebar.header("Input Details")

# User Inputs
subject_name = st.sidebar.text_input("Enter Subject Name")
syllabus = st.sidebar.text_area("Enter Syllabus (or upload)", height=200)
num_mcq = st.sidebar.slider("Number of MCQ questions", 0, 100, 10)
num_short = st.sidebar.slider("Number of Short Questions", 0, 100, 5)
num_long = st.sidebar.slider("Number of Long Questions", 0, 50, 3)

bl_level = st.sidebar.selectbox(
    "Select Bloom's Taxonomy Level",
    options=["Random (All Levels)", "BL1", "BL2", "BL3", "BL4", "BL5", "BL6"]
)

# Buttons with icons and tooltips
mcq_button = st.sidebar.button("Generate MCQs", help="Generate multiple choice questions based on your inputs")
short_button = st.sidebar.button("Generate Short Questions", help="Generate short answer questions")
long_button = st.sidebar.button("Generate Long Questions", help="Generate long answer questions")
all_button = st.sidebar.button("Generate All Questions", help="Generate all questions (MCQs, Short, and Long)")

# Display generated questions
if mcq_button:
    if subject_name and syllabus:
        with st.spinner('Generating MCQs...'):
            st.session_state.mcq_questions = generate_mcq_questions(subject_name, syllabus, num_mcq, bl_level)
            st.subheader("Generated MCQ Questions")
            st.markdown(f"**MCQ Questions (Level: {bl_level})**")
            st.text_area("Generated MCQ Questions", value=st.session_state.mcq_questions, height=300)

if short_button:
    if subject_name and syllabus:
        with st.spinner('Generating Short Questions...'):
            st.session_state.short_questions = generate_short_questions(subject_name, syllabus, num_short, bl_level)
            st.subheader("Generated Short Questions")
            st.markdown(f"**Short Answer Questions (Level: {bl_level})**")
            st.text_area("Generated Short Questions", value=st.session_state.short_questions, height=300)

if long_button:
    if subject_name and syllabus:
        with st.spinner('Generating Long Questions...'):
            st.session_state.long_questions = generate_long_questions(subject_name, syllabus, num_long, bl_level)
            st.subheader("Generated Long Questions")
            st.markdown(f"**Long Answer Questions (Level: {bl_level})**")
            st.text_area("Generated Long Questions", value=st.session_state.long_questions, height=300)

# Generate all questions (MCQs, Short, and Long)
if all_button:
    if subject_name and syllabus:
        with st.spinner('Generating All Questions...'):
            st.session_state.mcq_questions = generate_mcq_questions(subject_name, syllabus, num_mcq, bl_level)
            st.session_state.short_questions = generate_short_questions(subject_name, syllabus, num_short, bl_level)
            st.session_state.long_questions = generate_long_questions(subject_name, syllabus, num_long, bl_level)
            
            st.subheader("Generated All Questions")
            st.markdown(f"**MCQ Questions (Level: {bl_level})**")
            st.text_area("Generated MCQ Questions", value=st.session_state.mcq_questions, height=300)
            st.markdown(f"**Short Answer Questions (Level: {bl_level})**")
            st.text_area("Generated Short Questions", value=st.session_state.short_questions, height=300)
            st.markdown(f"**Long Answer Questions (Level: {bl_level})**")
            st.text_area("Generated Long Questions", value=st.session_state.long_questions, height=300)

# Download generated questions as DOCX
if st.sidebar.button("Generate All Questions as DOCX"):
    if st.session_state.mcq_questions or st.session_state.short_questions or st.session_state.long_questions:
        doc = docx.Document()
        if st.session_state.mcq_questions:
            doc.add_heading("Generated MCQ Questions", level=1)
            doc.add_paragraph(st.session_state.mcq_questions)
        if st.session_state.short_questions:
            doc.add_heading("Generated Short Answer Questions", level=1)
            doc.add_paragraph(st.session_state.short_questions)
        if st.session_state.long_questions:
            doc.add_heading("Generated Long Answer Questions", level=1)
            doc.add_paragraph(st.session_state.long_questions)

        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)

        st.download_button(
            label="Download as DOCX",
            data=doc_io,
            file_name="generated_questions.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# Add a footer with information or tips
st.markdown("<hr>", unsafe_allow_html=True)
st.markdown("### Tips: \n - Enter the subject name and syllabus to generate questions. \n - Choose the number of questions and Bloom's level.", unsafe_allow_html=True)
